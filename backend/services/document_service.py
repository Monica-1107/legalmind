import os
import PyPDF2
from docx import Document as DocxDocument
from datetime import datetime
from models.document import Document
from database import documents, db

def process_document(file_path, file_type):
    """
    Process an uploaded document to extract text and metadata
    
    Args:
        file_path (str): Path to the uploaded file
        file_type (str): MIME type of the file
        
    Returns:
        str: Document ID
    """
    content = ""
    metadata = {}
    title = os.path.basename(file_path)  # Default title is the filename
    
    try:
        # Extract content based on file type
        if file_type == 'application/pdf':
            content, pdf_metadata = extract_pdf_content(file_path)
            metadata.update(pdf_metadata)
            # Try to get a better title from PDF metadata
            if 'title' in pdf_metadata and pdf_metadata['title']:
                title = pdf_metadata['title']
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            content, doc_metadata = extract_docx_content(file_path)
            metadata.update(doc_metadata)
            # Try to get a better title from DOCX metadata
            if 'title' in doc_metadata and doc_metadata['title']:
                title = doc_metadata['title']
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Add basic metadata
        metadata["page_count"] = len(content.split('\n\n'))
        
        # Create document
        document = Document(
            filename=os.path.basename(file_path),
            file_path=file_path,
            file_type=file_type,
            content=content,
            metadata=metadata,
            title=title
        )
        
        # Insert into MongoDB
        result = documents.insert_one(document.to_dict())
        
        return str(result.inserted_id)
        
    except Exception as e:
        raise Exception(f"Error processing document: {str(e)}")

def extract_pdf_content(file_path):
    """Extract text and metadata from a PDF file"""
    content = ""
    metadata = {}
    
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        
        metadata["page_count"] = len(pdf_reader.pages)
        
        # Extract metadata
        if pdf_reader.metadata:
            info = pdf_reader.metadata
            metadata["author"] = info.author if hasattr(info, 'author') else None
            metadata["creator"] = info.creator if hasattr(info, 'creator') else None
            metadata["producer"] = info.producer if hasattr(info, 'producer') else None
            metadata["subject"] = info.subject if hasattr(info, 'subject') else None
            metadata["title"] = info.title if hasattr(info, 'title') else None
        
        # Extract text from each page
        for page in pdf_reader.pages:
            content += page.extract_text() + "\n\n"
    
    return content.strip(), metadata

def extract_docx_content(file_path):
    """Extract text and metadata from a DOCX file"""
    content = ""
    metadata = {}
    
    doc = DocxDocument(file_path)
    
    metadata["page_count"] = len(doc.sections)
    metadata["paragraph_count"] = len(doc.paragraphs)
    
    # Extract metadata
    core_props = doc.core_properties
    if core_props:
        metadata["author"] = core_props.author
        metadata["created"] = core_props.created.isoformat() if core_props.created else None
        metadata["modified"] = core_props.modified.isoformat() if core_props.modified else None
        metadata["title"] = core_props.title
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text:
            content += paragraph.text + "\n\n"
    
    return content.strip(), metadata